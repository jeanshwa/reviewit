import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
from pathlib import Path
import pypdf

# API clients
try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

try:
    import anthropic
except ImportError:
    anthropic = None

try:
    import google.generativeai as genai
except ImportError:
    genai = None

st.set_page_config(page_title="NSW Legal Analyzer Pro", page_icon="📜", layout="wide")

# ============== 密码保护 ==============
st.title("🔒 NSW Legal Analyzer Pro")

CORRECT_PASSWORD = st.secrets.get("APP_PASSWORD", "hello")
password = st.text_input("请输入访问密码", type="password")

if password != CORRECT_PASSWORD:
    st.warning("🔒 请输入正确密码")
    st.stop()

st.success("✅ 验证通过")

st.subheader("Contract of Sale & Lease 智能双语分析")

# ============== SIDEBAR ==============
with st.sidebar:
    st.header("API 设置")
    llm_provider = st.selectbox("选择 LLM", ["Grok (XAI)", "Gemini (Google)", "Claude (Anthropic)"])
    
    if llm_provider == "Grok (XAI)":
        api_key = st.secrets.get("XAI_API_KEY", st.text_input("XAI_API_KEY", type="password"))
    elif llm_provider == "Gemini (Google)":
        api_key = st.secrets.get("GOOGLE_API_KEY", st.text_input("GOOGLE_API_KEY", type="password"))
    else:
        api_key = st.secrets.get("ANTHROPIC_API_KEY", st.text_input("ANTHROPIC_API_KEY", type="password"))
    
    doc_type = st.selectbox("文档类型", ["Contract of Sale", "Lease"])

# ============== 文件上传 ==============
uploaded_file = st.file_uploader("📤 上传 NSW 合同或租赁文件", type=["pdf", "txt", "docx"])

def extract_text(uploaded_file):
    if uploaded_file.type == "application/pdf":
        try:
            reader = pypdf.PdfReader(uploaded_file)
            return "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
        except:
            return ""
    else:
        return uploaded_file.read().decode("utf-8", errors="ignore")

def get_prompt(text, doc_type):
    return f"""You are an expert NSW property solicitor specializing in conveyancing.

Focus **heavily** on **Special Conditions**, especially those that **modify or override the standard General Conditions** (e.g. cl.7 claims, cl.10 restrictions, risk passing, completion, disclosures).

Document Type: {doc_type}

Document Content:
{text[:15000]}

Provide detailed structured analysis in BOTH English and Chinese.

**ENGLISH ANALYSIS**
Executive Summary:
**Special Conditions Deep Analysis** (重点):
- Modifications to General Conditions:
- Key Risks (High/Medium/Low):
- Negotiation Recommendations:

**中文分析**
执行摘要：
**特别条款深度审查**（重点）：
- 对通用条款的修改：
- 主要风险：
- 谈判建议：
"""

def call_grok(api_key, prompt):
    if not OpenAI: return "Error: openai not installed"
    try:
        client = OpenAI(api_key=api_key, base_url="https://api.x.ai/v1")
        resp = client.chat.completions.create(model="grok-3", messages=[{"role":"user","content":prompt}], temperature=0.3, max_tokens=4000)
        return resp.choices[0].message.content
    except Exception as e: return f"Grok Error: {str(e)}"

def call_gemini(api_key, prompt):
    if not genai: return "Error: google-generativeai not installed"
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-2.0-flash")
        return model.generate_content(prompt).text
    except Exception as e: return f"Gemini Error: {str(e)}"

def call_anthropic(api_key, prompt):
    if not anthropic: return "Error: anthropic not installed"
    try:
        client = anthropic.Anthropic(api_key=api_key)
        resp = client.messages.create(model="claude-3-5-sonnet-20241022", max_tokens=4000, messages=[{"role":"user","content":prompt}])
        return resp.content[0].text
    except Exception as e: return f"Anthropic Error: {str(e)}"

if uploaded_file:
    text = extract_text(uploaded_file)
    st.success(f"✅ 上传成功: {uploaded_file.name}")

    if st.button("🚀 开始分析并生成双语报告", type="primary", use_container_width=True):
        if not api_key:
            st.error("请配置 API Key")
        else:
            with st.spinner(f"调用 {llm_provider} 分析中..."):
                prompt = get_prompt(text, doc_type)
                
                if llm_provider == "Grok (XAI)":
                    analysis = call_grok(api_key, prompt)
                elif llm_provider == "Gemini (Google)":
                    analysis = call_gemini(api_key, prompt)
                else:
                    analysis = call_anthropic(api_key, prompt)
                
                doc = Document()
                doc.add_heading('NSW Legal Analysis Report', 0)
                doc.add_paragraph(f"文件: {uploaded_file.name} | LLM: {llm_provider}")
                doc.add_paragraph(analysis)
                
                report_path = f"{Path(uploaded_file.name).stem}_NSW_Analysis_Report.docx"
                doc.save(report_path)
                
                with open(report_path, "rb") as f:
                    st.download_button("📥 下载中英双语报告", f, report_path)
                
                st.success("报告生成完成！")
                with st.expander("查看完整分析"):
                    st.markdown(analysis)

st.caption("完整版 | Special Conditions 重点审查 | Secrets 配置已支持")
